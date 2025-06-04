from fastapi import FastAPI, HTTPException # type: ignore
from fastapi.responses import FileResponse
from pydantic import BaseModel # type: ignore
from typing import List, Optional
import os
import groq  # type: ignore
from dotenv import load_dotenv  # type: ignore
from PyPDF2 import PdfReader # type: ignore
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml.ns import qn
import re

load_dotenv()
api_key = os.getenv("GROQ_API_KEY")
client = groq.Groq(api_key=api_key)

app = FastAPI()

# FUNCTION FOR GIVING API CALL TO GROQ
def generate_legal_document(prompt):
    """Handles the API call to generate legal documents."""
    response = client.chat.completions.create(
        model="gemma2-9b-it",
        max_completion_tokens=1024,
        temperature=1,
        stop=None,
        messages=[{"role": "system", "content": prompt}]
        
    )
    return response.choices[0].message.content

# Function to extract summarized text from PDF
def extract_text_from_pdf(pdf_path, max_pages=10, max_chars=5000):
    """Extracts limited text from the PDF to fit within token limits."""
    reader = PdfReader(pdf_path)
    extracted_text = []
    
    for i, page in enumerate(reader.pages[:max_pages]):  # Limit the number of pages
        text = page.extract_text()
        if text:
            extracted_text.append(text)
    
    full_text = "\n".join(extracted_text)
    
    # Truncate text if it exceeds the character limit
    return full_text[:max_chars] 

# Request models
class WritPetitionRequest(BaseModel):
    case_number: str
    year: int
    petitioner: str
    respondent: str
    court_name: str
    jurisdiction: str
    legal_grounds: str
    relief_sought: str
    benefits_offered: str
    supporting_documents: List[str]

class AffidavitRequest(BaseModel):
    case_number: str
    year: int
    petitioner: str
    respondent: str
    affiant_name: str
    order_date: str
    property_name: str
    statement_of_facts: str

class PatentApplicationRequest(BaseModel):
    application_number: str
    abstract: str
    year: int
    inventor_name: str
    assignee: str
    title: str
    field_of_invention: str
    background: str
    summary: str
    claims: str
    drawings_description: str

class AnnexureRequest(BaseModel):
    case_number: str
    year: int
    petitioner: str
    respondent: str
    annexure_title: str
    annexure_number: str
    description: str
    supporting_documents: List[str]

class WitnessStatementRequest(BaseModel):
    case_number: str
    year: int
    court_name: str
    witness_name: str
    witness_details: str
    statement: str

class ExhibitRequest(BaseModel):
    case_number: str
    year: int
    exhibit_number: str
    exhibit_title: str
    description: str
    attached_documents: List[str]

class ForensicReportRequest(BaseModel):
    case_number: str
    year: int
    forensic_expert: str
    forensic_field: str
    report_summary: str
    findings: str
    conclusion: str

class ExpertOpinionRequest(BaseModel):
    case_number: str
    year: int
    expert_name: str
    field_of_expertise: str
    opinion_summary: str
    detailed_opinion: str
    supporting_references: List[str]


# PROMPTS FOR GENERATION OF CORE LEGAL DOCS
@app.post("/generate/writ_petition")
def generate_writ_petition(data: WritPetitionRequest):
    context_text = extract_text_from_pdf(r"D:\ipd\SampleDocuments\Writ_petition.pdf")

    # Prompt Stuffing Template
    prompt = f"""
        You are a legal assistant tasked with generating a formal **Writ Petition** affidavit based on the given structure and input details. Maintain clarity, legal precision, and professional tone throughout.

        ---

        ### 📄 Context:  
        Refer to the following extracted format for guidance:

        {context_text}

        ---

        ### 📌 Input Parameters:
        - Case Number: {data.case_number} of {data.year}  
        - Petitioner: {data.petitioner}  
        - Respondent: {data.respondent}  
        - Court: {data.court_name}  
        - Jurisdiction: {data.jurisdiction}  
        - Legal Grounds: {data.legal_grounds}  
        - Relief Sought: {data.relief_sought}  
        - Supporting Documents: {', '.join(data.supporting_documents)}

        ---

        ### 🧾 Output Format:

        1. **Title & Case Details**  
        AFFIDAVIT  
        IN THE MATTER OF:  
        Case No. {data.case_number} of {data.year}  
        Petitioner: {data.petitioner}  
        Respondent: {data.respondent}  

        2. **Introduction (Details of the Affiant)**  
        - Full name, age, occupation, and address of the affiant.  
        - Relation to the case and purpose of submitting the affidavit.
        - If the introduction is not provided, omit this section.
        - If the introduction is provided, ensure it is clearly stated and formatted correctly.
        - If the introduction is in the form of a table, ensure it is clearly stated and formatted correctly.
        - If the introduction is in the form of a list, ensure it is clearly stated and formatted correctly.
        - If the introduction is in the form of a paragraph, ensure it is clearly stated and formatted correctly.

        3. **Statement of Facts**  
        - Present all relevant facts in numbered points.  
        - Maintain chronological order and legal clarity.
        - If the statement of facts is not provided, omit this section.
        - If the statement of facts is provided, ensure it is clearly stated and formatted correctly.
        - If the statement of facts is in the form of a table, ensure it is clearly stated and formatted correctly.
        - If the statement of facts is in the form of a list, ensure it is clearly stated and formatted correctly.
        - If the statement of facts is in the form of a paragraph, ensure it is clearly stated and formatted correctly.


        4. **Benefits Offered**  
        - {data.benefits_offered}
        - Structure the benefits offered in a proper legal format.
        - If no benefits are offered, omit this section.
        - If benefits are offered, ensure they are clearly stated and formatted correctly.
        - If the benefits are in the form of rent and corpus, ensure they are clearly stated and formatted correctly.
        - If the benefits are in the form of any other property, ensure they are clearly stated and formatted correctly.
        - If the benefits are in the form of any other asset, ensure they are clearly stated and formatted correctly.
        - If the benefits are in the form of any other liability, ensure they are clearly stated and formatted correctly.
        - If the benefits are in the form of any other obligation, ensure they are clearly stated and formatted correctly.
        - If the benefits are in the form of any other right, ensure they are clearly stated and formatted correctly.
        - If the benefits are in the form of any other interest, ensure they are clearly stated and formatted correctly.
        - If the benefits are in the form of any other claim, ensure they are clearly stated and formatted correctly.
        - If the benefits are in the form of any other demand, ensure they are clearly stated and formatted correctly.


        5. **Compliance with Regulations**  
        - Declare compliance with all relevant rules, acts, or orders concerning the case/property.
        - If the compliance with regulations is not provided, omit this section.
        - If the compliance with regulations is provided, ensure it is clearly stated and formatted correctly.
        - If the compliance with regulations is in the form of a table, ensure it is clearly stated and formatted correctly.
        - If the compliance with regulations is in the form of a list, ensure it is clearly stated and formatted correctly.
        - If the compliance with regulations is in the form of a paragraph, ensure it is clearly stated and formatted correctly.

        6. **Oath & Affirmation**  
        - A solemn affirmation that all information provided is true and correct to the best of the affiant's knowledge and belief.
        - If the oath and affirmation is not provided, omit this section.
        - If the oath and affirmation is provided, ensure it is clearly stated and formatted correctly.
        - If the oath and affirmation is in the form of a table, ensure it is clearly stated and formatted correctly.
        - If the oath and affirmation is in the form of a list, ensure it is clearly stated and formatted correctly.
        - If the oath and affirmation is in the form of a paragraph, ensure it is clearly stated and formatted correctly.

        7. **Signature & Notarization**  
        - Signature of the affiant with date and location.  
        - Notary section for official seal and authentication.
        - If the signature and notarization is not provided, omit this section.
        - If the signature and notarization is provided, ensure it is clearly stated and formatted correctly.
        - If the signature and notarization is in the form of a table, ensure it is clearly stated and formatted correctly.
        - If the signature and notarization is in the form of a list, ensure it is clearly stated and formatted correctly.
        - If the signature and notarization is in the form of a paragraph, ensure it is clearly stated and formatted correctly.

        ---

        ### ⚖️ Instructions:
        - Use a formal, structured legal tone.
        - Automatically omit optional fields if data is missing.
        - Ensure all sections are well-formatted and professionally worded.
        - Structure the document in proper legal format with numbered clauses where necessary, ensuring clarity and legal precision.
        - Ensure that the document is structured with numbered clauses where necessary, ensuring clarity and legal precision.

    """
    safe_case_number = data.case_number.replace("/", "_")
    filename = f"WritPetition_{safe_case_number}_{data.year}.docx"

    full_path = save_to_docx(
        content=generate_legal_document(prompt),
        filename=filename,
        case_number=data.case_number,
        year=data.year
    )

    return FileResponse(
        path=full_path,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        filename=filename
    )




@app.post("/generate/affidavit")
def generate_affidavit(data: AffidavitRequest):
    context_text = extract_text_from_pdf(r"D:\ipd\SampleDocuments\Affidavit.pdf")

    # Prompt Stuffing Template
    prompt = f"""
        You are a legal assistant tasked with generating a formal **Affidavit** in a professional legal tone based on the details provided. Maintain structure, clarity, and precision, adhering to standard affidavit formats.

        ---

        ### 📄 Context Reference:
        Please refer to the following example for layout guidance:
        {context_text}

        ---

        ### 📌 Input Details:
        - Case Number: {data.case_number} of {data.year}
        - Petitioner: {data.petitioner}
        - Respondent: {data.respondent}
        - Affiant Name: {data.affiant_name}
        - Order Date: {data.order_date}
        - Property Name: {data.property_name}
        - Statement of Facts: {data.statement_of_facts}

        ---

        ### 🧾 Output Format:

        1. **Title & Case Details**  
        AFFIDAVIT  
        IN THE MATTER OF:  
        Case No. {data.case_number} of {data.year}  
        Petitioner: {data.petitioner}  
        Respondent: {data.respondent}  

        2. **Introduction (Details of the Affiant)**  
        - Name: {data.affiant_name}  
        - Role in case: Mention the capacity in which affiant is filing the affidavit.  
        - Date of the order: {data.order_date}  
        - Property under consideration: {data.property_name} 
         

        3. **Statement of Facts**  
        - Present all relevant facts in bullet points or numbered clauses.
        - Ensure chronological flow and clear legal articulation.  
        - Input facts: {data.statement_of_facts}
        - If the statement of facts is not provided, omit this section.
        - If the statement of facts is provided, ensure it is clearly stated and formatted correctly.
        - If the statement of facts is in the form of a table, ensure it is clearly stated and formatted correctly.
        - If the statement of facts is in the form of a list, ensure it is clearly stated and formatted correctly.
        - If the statement of facts is in the form of a paragraph, ensure it is clearly stated and formatted correctly.

        4. **Benefits Offered (if applicable)**  
        - Mention any benefits like rent, corpus, or other entitlements.
        - If not applicable, this section should be omitted.
        - If the benefits are in the form of rent and corpus, ensure they are clearly stated and formatted correctly.
        - If the benefits are in the form of any other property, ensure they are clearly stated and formatted correctly.
        - If the benefits are in the form of any other asset, ensure they are clearly stated and formatted correctly.
        - If the benefits are in the form of any other liability, ensure they are clearly stated and formatted correctly.
        - If the benefits are in the form of any other obligation, ensure they are clearly stated and formatted correctly.
        - If the benefits are in the form of any other right, ensure they are clearly stated and formatted correctly.

        5. **Compliance with Regulations**  
        - Declare adherence to applicable laws, orders, or government norms related to the matter.
        - Use precise legal language.
        - If the compliance with regulations is not provided, omit this section.
        - If the compliance with regulations is provided, ensure it is clearly stated and formatted correctly.
        - If the compliance with regulations is in the form of a table, ensure it is clearly stated and formatted correctly.
        - If the compliance with regulations is in the form of a list, ensure it is clearly stated and formatted correctly.
        - If the compliance with regulations is in the form of a paragraph, ensure it is clearly stated and formatted correctly.

        6. **Oath & Affirmation**  
        - The affiant solemnly affirms that the statements made above are true and correct to the best of their knowledge and belief.
        - If the oath and affirmation is not provided, omit this section.
        - If the oath and affirmation is provided, ensure it is clearly stated and formatted correctly.
        
        7. **Signature & Notarization**  
        - Signature of the affiant with date and place.
        - Include space for notarization and official seal.

        ---

        ### ⚖️ Legal Tone & Instructions:
        - Maintain a formal legal tone and precise language throughout.
        - Automatically skip sections where input is missing or marked as not applicable.
        - Format all sections consistently using bullet points or numbered lists for readability.
        - Ensure proper legal formatting, indentation, and spacing.
        - Follow the reference format strictly but adapt it to reflect the current case details.

    """
    safe_case_number = data.case_number.replace("/", "_")
    filename = f"Affidavit_{safe_case_number}_{data.year}.docx"

    full_path = save_to_docx(
        content=generate_legal_document(prompt),
        filename=filename,
        case_number=data.case_number,
        year=data.year
    )

    return FileResponse(
        path=full_path,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        filename=filename
    )




@app.post("/generate/patent_application")
def generate_patent_application(data: PatentApplicationRequest):
    context_text = extract_text_from_pdf(r"D:\ipd\SampleDocuments\Patent_application.pdf")

    # Prompt Stuffing Template
    prompt = f"""
        You are a legal documentation assistant. Generate a **Patent Application** in a structured and legally precise format based on the input data provided below. Maintain professional tone, clarity, and consistency with standard patent filing formats.

        ---

        ### 📄 Context Reference:
        Below is the extracted structure from a previously filed patent application for format guidance:
        {context_text}

        ---

        ### 📌 Input Details:
        - Application Number: {data.application_number} of {data.year}
        - Inventor Name: {data.inventor_name}
        - Abstract: {data.abstract}
        - Assignee: {data.assignee}
        - Title of Invention: {data.title}
        - Field of Invention: {data.field_of_invention}
        - Background: {data.background}
        - Summary: {data.summary}
        - Claims: {data.claims}
        - Description of Drawings: {data.drawings_description}

        ---

        ### 🧾 Output Format:

        1. **Title & Application Details**  
        - Title of Invention: {data.title}  
        - Application No.: {data.application_number} of {data.year}  
        - Inventor: {data.inventor_name}  
        - Assignee: {data.assignee}  

        2. **Field of the Invention**  
        - {data.field_of_invention}  
        - Clearly state the technical domain and scope of the invention.  
        - If field of invention is not provided, omit this section.
        - If field of invention is provided, ensure it is clearly stated and formatted correctly.
        - If field of invention is in the form of a table, ensure it is clearly stated and formatted correctly.
        - If field of invention is in the form of a list, ensure it is clearly stated and formatted correctly.
        - If field of invention is in the form of a paragraph, ensure it is clearly stated and formatted correctly.

        3. **Background of the Invention**  
        - {data.background}  
        - Include prior art and the problem the invention seeks to solve.  
        - If background is not provided, omit this section.
        - If background is provided, ensure it is clearly stated and formatted correctly.
        - If background is in the form of a table, ensure it is clearly stated and formatted correctly.
        - If background is in the form of a list, ensure it is clearly stated and formatted correctly.
        - If background is in the form of a paragraph, ensure it is clearly stated and formatted correctly.

        4. **Summary of the Invention**  
        - {data.summary}  
        - Summarize the novel features and advantages of the invention.  
        - If summary is not provided, omit this section.
        - If summary is provided, ensure it is clearly stated and formatted correctly.
        - If summary is in the form of a table, ensure it is clearly stated and formatted correctly.
        - If summary is in the form of a list, ensure it is clearly stated and formatted correctly.
        - If summary is in the form of a paragraph, ensure it is clearly stated and formatted correctly.

        5. **Detailed Description of the Invention**  
        - Expand on the technical aspects of the invention and how it works.  
        - Include components, processes, or systems if applicable.  
        - If omitted in input, skip this section.
        - If detailed description of the invention is not provided, omit this section.
        - If detailed description of the invention is provided, ensure it is clearly stated and formatted correctly.
        - If detailed description of the invention is in the form of a table, ensure it is clearly stated and formatted correctly.
        - If detailed description of the invention is in the form of a list, ensure it is clearly stated and formatted correctly.
        - If detailed description of the invention is in the form of a paragraph, ensure it is clearly stated and formatted correctly.

        6. **Claims**  
        - Define the scope of legal protection sought.  
        - Use numbered claims, each precisely stating a distinct element or feature.  
        - {data.claims}  
        - Ensure clarity, specificity, and legal enforceability of each claim.  
        - If claims are not provided, omit this section.
        - If claims are provided, ensure it is clearly stated and formatted correctly.
        - If claims are in the form of a table, ensure it is clearly stated and formatted correctly.
        - If claims are in the form of a list, ensure it is clearly stated and formatted correctly.
        - If claims are in the form of a paragraph, ensure it is clearly stated and formatted correctly.

        7. **Brief Description of Drawings**  
        - {data.drawings_description}  
        - If applicable, provide brief narrative descriptions of accompanying figures.  
        - If drawings are not available, omit this section.
        - If drawings are provided, ensure it is clearly stated and formatted correctly.
        - If drawings are in the form of a table, ensure it is clearly stated and formatted correctly.
        - If drawings are in the form of a list, ensure it is clearly stated and formatted correctly.
        - If drawings are in the form of a paragraph, ensure it is clearly stated and formatted correctly.

        8. **Abstract**  
        - {data.abstract}
        - Provide a concise summary of the invention (max 150 words).  
        - This should be generated based on the title, summary, and claims.  
        - Ensure it is self-contained and informative for a general reader.
        - If abstract is not provided, omit this section.
        - If abstract is provided, ensure it is clearly stated and formatted correctly.
        - If abstract is in the form of a table, ensure it is clearly stated and formatted correctly.
        - If abstract is in the form of a list, ensure it is clearly stated and formatted correctly.
        - If abstract is in the form of a paragraph, ensure it is clearly stated and formatted correctly.

        ---

        ### ⚖️ Drafting Instructions:
        - Ensure formal, technical, and legally sound tone.
        - Automatically omit any optional sections if corresponding data is missing.
        - Number all claims and relevant sub-sections for easy reference.
        - Maintain professional formatting, with bullet points and headers for clarity.
        - The document should reflect the structure of standard patent filings.

    """
    safe_case_number = data.application_number.replace("/", "_")
    filename = f"PatentApplication_{safe_case_number}_{data.year}.docx"

    full_path = save_to_docx(
        content=generate_legal_document(prompt),
        filename=filename,
        case_number=data.application_number,
        year=data.year
    )

    return FileResponse(
        path=full_path,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        filename=filename
    )




@app.post("/generate/annexure")
def generate_annexure(data: AnnexureRequest):
    context_text = extract_text_from_pdf(r"D:\ipd\SampleDocuments\Annexure.pdf")

    # Prompt Stuffing Template
    prompt = f"""
        You are a legal assistant responsible for drafting a formal **Annexure** for a court case. The document should follow legal formatting standards and include all necessary sections as per the input provided.

        ---

        ### 📄 Contextual Reference:
        The following text has been extracted from a previously approved annexure and should be used as a format guide:
        {context_text}

        ---

        ### 📌 Input Details:
        - Case Number: {data.case_number} of {data.year}
        - Petitioner: {data.petitioner}
        - Respondent: {data.respondent}
        - Annexure Title: {data.annexure_title}
        - Annexure Number: {data.annexure_number}
        - Description: {data.description}
        - Supporting Documents: {', '.join(data.supporting_documents)}

        ---

        ### 🧾 Output Format:

        1. **Title Section**  
        - ANNEXURE {data.annexure_number}  
        - IN THE MATTER OF: Case No. {data.case_number} of {data.year}  
        - Petitioner: {data.petitioner}  
        - Respondent: {data.respondent}

        2. **Introduction**  
        - Briefly describe the purpose and relevance of this annexure.  
        - Reference the title: "{data.annexure_title}"  
        - Clearly establish the context in relation to the primary case.

        3. **Detailed Description**  
        - Elaborate on the content of the annexure.  
        - Include factual background, if applicable.  
        - {data.description}  
        - Use clear, numbered paragraphs if content is extensive.
        

        4. **List of Supporting Documents**  
        - Enumerate the documents referenced or enclosed with this annexure.  
        - Format as a numbered list:  
        {"".join([f"    {i+1}. {doc}\n" for i, doc in enumerate(data.supporting_documents)]) if data.supporting_documents else "    (None Provided)"}

        5. **Affirmation**  
        - Declare the authenticity of the contents and the supporting documents enclosed.  
        - Include a professional affirmation in legal language.  
        - Example: "I hereby affirm that the contents of this annexure and the documents appended herewith are true and correct to the best of my knowledge and belief."

        6. **Date & Signature**  
        - Include signature block with affiant/legal representative's name.  
        - Mention location and date.  
        - Add space for official seal if required.

        ---

        ### ⚖️ Formatting & Legal Drafting Instructions:
        - Use a formal, structured legal tone.
        - Automatically omit any optional sections if data is missing.
        - Number relevant sections and documents clearly.
        - Ensure clarity, coherence, and precision in wording.
        - Match standard Indian legal documentation practices.

    """
    safe_case_number = data.annexure_number.replace("/", "_")
    filename = f"Annexure_{safe_case_number}_{data.year}.docx"

    full_path = save_to_docx(
        content=generate_legal_document(prompt),
        filename=filename,
        case_number=data.annexure_number,
        year=data.year
    )

    return FileResponse(
        path=full_path,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        filename=filename
    )




@app.post("/generate/witness_statement")
def generate_witness_statement(data: WitnessStatementRequest):
    context_text = extract_text_from_pdf("D:/Paralegal_App/backend/SampleDocuments/Witness_statement.pdf")

    # Prompt Stuffing Template
    prompt = f"""
        You are a legal assistant tasked with drafting a formal **Witness Statement** for submission in a court of law. Use a professional and structured legal tone, ensuring clarity, factual accuracy, and legal admissibility.

        ---

        ### 📄 Format Reference:
        Use the following extracted legal template for formatting and tone:
        {context_text}

        ---

        ### 📌 Input Parameters:
        - Case Number: {data.case_number} of {data.year}
        - Court: {data.court_name}
        - Witness Name: {data.witness_name}
        - Witness Details: {data.witness_details}
        - Statement: {data.statement}

        ---

        ### 🧾 Output Structure:

        1. **Title & Case Details**  
        - WITNESS STATEMENT  
        - IN THE MATTER OF: Case No. {data.case_number} of {data.year}  
        - Court: {data.court_name}

        2. **Witness Introduction**  
        - Name: {data.witness_name}  
        - {data.witness_details}  
        - Clearly state the age, occupation, and relationship of the witness to the case.

        3. **Statement of Facts**  
        - A factual, chronological account of what the witness observed or experienced.  
        - Format in numbered paragraphs for legal readability.  
        - Ensure each point is distinct, logically sequenced, and free from ambiguity.  
        - {data.statement}

        4. **Affirmation**  
        - The witness shall affirm as follows:  
        > "I, {data.witness_name}, do hereby solemnly affirm and declare that the foregoing contents of this statement are true and correct to the best of my knowledge and belief. I understand that giving false evidence is a punishable offence under the law."

        5. **Signature & Date**  
        - Signature of the witness  
        - Date and location of signing  
        - Space for Notary or Court Officer if applicable

        ---

        ### ⚖️ Drafting Guidelines:
        - Maintain a formal, legal tone throughout the statement.
        - Ensure sections are clearly demarcated and logically ordered.
        - If any input field is missing or empty, omit that section automatically.
        - Follow best practices of witness affidavit admissibility under Indian Evidence Law.

    """
    safe_case_number = data.case_number.replace("/", "_")
    filename = f"WitnessStatement_{safe_case_number}_{data.year}.docx"

    full_path = save_to_docx(
        content=generate_legal_document(prompt),
        filename=filename,
        case_number=data.case_number,
        year=data.year
    )

    return FileResponse(
        path=full_path,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        filename=filename
    )




@app.post("/generate/exhibit")
def generate_exhibit(data: ExhibitRequest):
    context_text = extract_text_from_pdf("D:/Paralegal_App/backend/SampleDocuments/Exhibit.pdf")

    # Prompt Stuffing Template
    prompt = f"""
        You are a legal assistant tasked with generating a formal **Exhibit Document** for court submission. Use a structured legal format with appropriate headings, professional tone, and ensure clarity and admissibility in legal proceedings.

        ---

        ### 📄 Format Reference:
        Use the following extracted format for tone and structure:
        {context_text}

        ---

        ### 📌 Input Parameters:
        - Case Number: {data.case_number} of {data.year}
        - Exhibit Number: {data.exhibit_number}
        - Exhibit Title: {data.exhibit_title}
        - Description: {data.description}
        - Attached Documents: {', '.join(data.attached_documents)}

        ---

        ### 🧾 Output Structure:

        1. **Title & Case Details**  
        - EXHIBIT – {data.exhibit_number}  
        - Title: {data.exhibit_title}  
        - IN THE MATTER OF: Case No. {data.case_number} of {data.year}

        2. **Description of Exhibit**  
        - Clearly state the purpose and legal relevance of this exhibit.  
        - {data.description}  
        - Present facts objectively and concisely.

        3. **List of Attached Documents**  
        - Enumerate each document forming part of the exhibit.  
        - Format as a numbered list or bullet points.  
        - Attached Documents: {', '.join(data.attached_documents)}

        4. **Certification of Authenticity**  
        - Include a declaration such as:  
        > "I hereby certify that the contents of this exhibit and the documents annexed herewith are true, correct, and genuine to the best of my knowledge and belief."

        5. **Signature & Date**  
        - Signature of the submitting party or authorized legal representative  
        - Date and Place of submission  
        - Include a space for notarization if required

        ---

        ### ⚖️ Drafting Guidelines:
        - Use formal legal language throughout.  
        - Maintain proper formatting for readability and legal clarity.  
        - If any input is missing, automatically omit the related section.  
        - Ensure alignment with court exhibit submission standards.

    """
    safe_case_number = data.exhibit_number.replace("/", "_")
    filename = f"Exhibit_{safe_case_number}_{data.year}.docx"
    full_path = save_to_docx(
        content=generate_legal_document(prompt),
        filename=filename,
        case_number=data.exhibit_number,
        year=data.year
    )

    return FileResponse(
        path=full_path,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        filename=filename
    )



@app.post("/generate/forensic_report")
def generate_forensic_report(data: ForensicReportRequest):
    context_text = extract_text_from_pdf("D:/Paralegal_App/backend/SampleDocuments/Forensic_report.pdf")

    prompt = f"""
        You are a legal assistant tasked with generating a formal **Forensic Report** for judicial submission. The report should be written in a professional and structured legal format with technical clarity, ensuring factual accuracy and evidentiary compliance.

        ---

        ### 📄 Format Reference:
        Use the following extracted format as a template for tone and layout:
        {context_text}

        ---

        ### 📌 Input Parameters:
        - Case Number: {data.case_number} of {data.year}
        - Forensic Expert: {data.forensic_expert}
        - Field of Analysis: {data.forensic_field}
        - Report Summary: {data.report_summary}
        - Findings: {data.findings}
        - Conclusion: {data.conclusion}

        ---

        ### 🧾 Output Structure:

        1. **Title & Case Details**  
        - FORENSIC REPORT  
        - IN THE MATTER OF: Case No. {data.case_number} of {data.year}  
        - Field of Analysis: {data.forensic_field}

        2. **Expert Introduction**  
        - Name: {data.forensic_expert}  
        - Qualifications and professional background  
        - Area of specialization and experience relevant to this case

        3. **Summary of the Investigation**  
        - {data.report_summary}  
        - Include objective overview of materials examined, methods used, and purpose of investigation

        4. **Findings**  
        - {data.findings}  
        - Use bullet points or numbered format if multiple observations  
        - Include technical data, relevant measurements, visual/chemical/biometric analysis, etc.  
        - Ensure scientific clarity and traceability of results

        5. **Conclusion**  
        - {data.conclusion}  
        - Provide a professional opinion based on forensic evidence  
        - Mention limitations (if any), standard disclaimers, and basis for expert judgment

        6. **Signature & Certification**  
        - Full name and signature of the forensic expert  
        - Date and place of submission  
        - Include a declaration such as:  
        > "I hereby certify that this report is true and accurate to the best of my professional knowledge, and the findings herein are based on standard forensic procedures."

        ---

        ### ⚖️ Drafting Instructions:
        - Ensure legal and scientific precision  
        - Maintain structured formatting and formal tone  
        - If any section data is missing, omit it gracefully  
        - Ensure the report is suitable for court proceedings and expert cross-examination

    """
    safe_case_number = data.case_number.replace("/", "_")
    filename = f"ForensicReport_{safe_case_number}_{data.year}.docx"

    full_path = save_to_docx(
        content=generate_legal_document(prompt),
        filename=filename,
        case_number=data.case_number,
        year=data.year
    )

    return FileResponse(
        path=full_path,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        filename=filename
    )




@app.post("/generate/expert_opinion")
def generate_expert_opinion(data: ExpertOpinionRequest):
    context_text = extract_text_from_pdf("D:/Paralegal_App/backend/SampleDocuments/Expert_opinion.pdf")    
    prompt = f"""
        You are a legal assistant tasked with drafting a formal **Expert Opinion** for submission in legal proceedings. The document must follow professional legal standards, provide clear expert reasoning, and be admissible in court.

        ---

        ### 📄 Format Reference:
        Use the following extracted format as a template for tone, layout, and structure:
        {context_text}

        ---

        ### 📌 Input Parameters:
        - Case Number: {data.case_number} of {data.year}
        - Expert Name: {data.expert_name}
        - Field of Expertise: {data.field_of_expertise}
        - Opinion Summary: {data.opinion_summary}
        - Detailed Opinion: {data.detailed_opinion}
        - Supporting References: {', '.join(data.supporting_references)}

        ---

        ### 🧾 Output Format:

        1. **Title & Case Details**  
        - EXPERT OPINION  
        - Case No: {data.case_number} of {data.year}  
        - Field of Expertise: {data.field_of_expertise}

        2. **Expert Introduction**  
        - Name: {data.expert_name}  
        - Professional qualifications and certifications  
        - Area of specialization and years of experience  
        - Past notable cases (if applicable)

        3. **Summary of Opinion**  
        - {data.opinion_summary}  
        - Provide a brief overview of the expert's stance on the matter

        4. **Detailed Explanation & Analysis**  
        - {data.detailed_opinion}  
        - Present in clear, logically structured paragraphs or bullet points  
        - Include references to relevant scientific, technical, or legal principles  
        - Highlight methodology, assumptions, and rationale

        5. **Supporting References**  
        - Cite {', '.join(data.supporting_references)}  
        - Include any relevant case laws, legal precedents, journal articles, or technical standards  
        - If no references are available, omit this section

        6. **Conclusion**  
        - Restate the opinion based on presented facts and expertise  
        - Include standard disclaimers or limitations if applicable

        7. **Signature & Date**  
        - Full name and signature of the expert  
        - Date and location of the opinion draft  
        - Include declaration:  
        > "I hereby affirm that the opinion provided herein is true to the best of my knowledge and formed in accordance with my professional expertise."

        ---

        ### ⚖️ Legal Drafting Instructions:
        - Maintain a formal legal tone throughout  
        - Omit optional fields gracefully if data is not provided  
        - Ensure logical flow and legal readability  
        - Use numbered sections and subheadings where necessary for clarity and professionalism

    """
    safe_case_number = data.case_number.replace("/", "_")
    filename = f"ExpertOpinion_{safe_case_number}_{data.year}.docx"
    full_path = save_to_docx(
        content=generate_legal_document(prompt),
        filename=filename,
        case_number=data.case_number,
        year=data.year
    )

    return FileResponse(
        path=full_path,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        filename=filename
    )




def sanitize_filename(filename: str) -> str:
    return re.sub(r'[\\/*?:"<>|]', "_", filename)

def save_to_docx(content, filename, case_number=None, year=None, output_dir="D:/Paralegal_App/backend/Downloads"):
    """Save legal content to a styled .docx with legal formatting."""
    doc = Document()

    # Header
    if case_number and year:
        header = doc.sections[0].header
        para = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
        para.text = f"Case No: {case_number} | Year: {year}"
        para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        para.runs[0].font.size = Pt(10)
        para.runs[0].bold = True

    # Clean markdown from input
    lines = content.strip().splitlines()
    counter = 1
    for line in lines:
        line = line.strip()
        if not line:
            continue

        # Remove markdown artifacts
        line = re.sub(r"^#+\s*", "", line)
        line = line.replace("**", "")
        line = line.replace("*", "")
        line = line.replace("–", "-")

        # Section headers
        if re.match(r"^[A-Z\s&]+$", line) or line.endswith(":"):
            para = doc.add_paragraph()
            run = para.add_run(line.upper())
            run.bold = True
            run.font.size = Pt(12)
            para.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

        # Numbered paragraphs
        elif len(line) > 80:
            para = doc.add_paragraph()
            run = para.add_run(f"{counter}. {line}")
            run.font.size = Pt(11)
            para.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
            counter += 1

        # Short lines (e.g., sign-offs)
        else:
            para = doc.add_paragraph()
            run = para.add_run(line)
            run.font.size = Pt(11)
            para.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

    # Footer disclaimer
    section = doc.sections[0]
    footer = section.footer
    para = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
    para.text = "This is a system-generated legal document. Please verify with a licensed legal professional."
    para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    para.runs[0].font.size = Pt(9)

    # Save
    safe_name = sanitize_filename(filename)
    if not os.path.exists(output_dir):
        os.makedirs(output_dir)
    full_path = os.path.join(output_dir, safe_name)
    doc.save(full_path)

    print("Document saved to " + full_path)
    return full_path

